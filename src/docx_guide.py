import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger("DAIS_Intraday.DocxGuide")

def build_docx_guide(cfg: dict, outdir: str) -> str:
    """
    Generates a professional, dynamically synchronized MS Word operational manual.
    Automatically parses configuration parameters to ensure operational accuracy.
    """
    docx_filename = os.path.join(outdir, "DAIS_Intraday_Operational_Manual.docx")
    
    try:
        doc = Document()
        
        # Configure standard page setup layouts
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            
        # Core Palette Definitions (Slate Theme)
        COLOR_DARK = RGBColor(30, 41, 59)
        COLOR_SKY = RGBColor(2, 132, 199)
        COLOR_TEXT = RGBColor(71, 85, 105)

        # ------------------------------------------------------------
        # DOCUMENT TITLE & HEADER SYSTEM
        # ------------------------------------------------------------
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_p.add_run("DAIS Platform: Intraday Operational Manual")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = COLOR_DARK
        
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(
            f"System Deployment Ledger: {datetime.now().strftime('%Y-%m-%d')} | "
            f"Execution Granularity: {cfg.get('bar_interval', '5m')} Bars"
        )
        meta_run.font.name = 'Arial'
        meta_run.font.size = Pt(10)
        meta_run.font.italic = True
        meta_run.font.color.rgb = COLOR_TEXT
        
        doc.add_paragraph().add_run("1. Algorithmic Infrastructure Core Summary").font.size = Pt(16)
        doc.paragraphs[-1].runs[0].font.bold = True
        doc.paragraphs[-1].runs[0].font.color.rgb = COLOR_SKY
        
        intro_text = (
            "The Dynamic Asymmetric Inventory‑Aware Strategy (DAIS) platform operates as a "
            "high-frequency execution framework designed to capture short-term micro-momentum "
            "anomalies while preserving core underlying dividend distribution profiles. By using "
            "a vectorized tracking loop, the system dynamically scales position structures up or down "
            "based on real-time asset beta evaluations and directional moving average crossovers."
        )
        p = doc.add_paragraph(intro_text)
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(10.5)

        # ------------------------------------------------------------
        # DYNAMIC PARAMETER ANALYSIS MATRIX
        # ------------------------------------------------------------
        doc.add_paragraph().add_run("2. Active Production Variable Configuration").font.size = Pt(16)
        doc.paragraphs[-1].runs[0].font.bold = True
        doc.paragraphs[-1].runs[0].font.color.rgb = COLOR_SKY
        
        doc.add_paragraph("The following matrix displays the active parameter limits parsed directly from config.yaml:")
        
        # Build 2-column parameter configuration grid
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Configuration Key Name'
        hdr_cells[1].text = 'Active Parsed Runtime Limit Value'
        
        # Format table header cells
        for cell in hdr_cells:
            cell.width = Inches(3.25)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Apply programmatic background shading via XML manipulation
            from docx.oxml import parse_xml
            shading = parse_xml(r'<w:shd {} w:fill="1E293B"/>'.format('xmlns:w="http://openxmlformats.org"'))
            cell._tc.get_or_add_tcPr().append(shading)

        # Populate runtime keys from configuration layout dict dynamically
        param_keys = [
            'initial_capital', 'core_buy_amt', 'base_buy_shares', 
            'base_sell_shares', 'inventory_floor_shares', 'slippage_bps',
            'short_window', 'long_window', 'force_flat_at_close'
        ]
        
        for key in param_keys:
            if key in cfg or key.replace('_shares', '') in cfg:
                actual_key = key if key in cfg else key.replace('_shares', '')
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(cfg[actual_key])
                
                # Apply standard column widths
                row_cells[0].width = Inches(3.25)
                row_cells[1].width = Inches(3.25)
                
                # Format cell fonts for scannability
                for cell in row_cells:
                    cell.paragraphs[0].style.font.name = 'Arial'
                    cell.paragraphs[0].style.font.size = Pt(10)

        doc.add_paragraph() # Add standard spacer
        
        # ------------------------------------------------------------
        # PRODUCTION RISK ENGINE MANAGEMENT
        # ------------------------------------------------------------
        doc.add_paragraph().add_run("3. Risk Control Execution Directives").font.size = Pt(16)
        doc.paragraphs[-1].runs[0].font.bold = True
        doc.paragraphs[-1].runs[0].font.color.rgb = COLOR_SKY
        
        directives = [
            "Hard Trend Line Stop: Open orders are locked out if the executing ticker falls beneath its intraday moving average line.",
            "Cost Basis Tracking Shield: Sells are automatically rejected if the current market execution fill price drops below average position cost.",
            "Session Liquidation Window: If force_flat_at_close is set to True, the system liquidates open inventory blocks before the close to eliminate overnight pricing gaps."
        ]
        
        for directive in directives:
            doc.add_paragraph(directive, style='List Bullet')

        doc.save(docx_filename)
        logger.info(f"Dynamically synchronized Word operational manual generated at: {docx_filename}")
        return docx_filename
        
    except Exception as e:
        logger.error(f"Failed to compile document elements to Word manual layout: {e}", exc_info=True)
        raise e
